import docx

doc = docx.Document("D:/variant/2lab.docx")
new_doc = docx.Document()
N = []
Y = []
for paragraph in doc.paragraphs:
    shet = 0
    for sim in paragraph.text:
        if ord(sim) >= 33 and ord(sim) <= 64:
            shet += 1
    if shet == 0:
        Y.append(paragraph.text)
        print("ДА - ",paragraph.text)
    else:
        N.append(paragraph.text)
        print("НЕТ - ", paragraph.text)
print("Колличество вхождений в Y - ",len(Y))
print("Колличество вхождений в N - ",len(N))
