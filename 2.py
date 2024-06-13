# Кто хорошо трудится, тому есть чем хвалиться
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
import mtk2

def cript():
    doc = docx.Document("D:/variant/2lab.docx")
    new_doc = docx.Document()
    posl = str(input("Введите пословицу - "))
    bys = mtk2.MTK2_code(posl)
    print("Пословица в двоичной последовательности ", mtk2.MTK2_code(posl))
    length = 0
               
    for paragraph in doc.paragraphs:
        par = new_doc.add_paragraph()
        for sim in paragraph.text:
            if length < len(bys):
                if bys[length] == "1":
                    par.add_run(sim).font.color.rgb = RGBColor(1,0,1)
                if bys[length] == "0":
                    par.add_run(sim).font.color.rgb = RGBColor(0,0,0)
                length += 1
            else:
                par.add_run(sim).font.color.rgb = RGBColor(0,0,0)

    new_doc.save("D:/variant/2_lab.docx")

def encript():
    doc = docx.Document("D:/variant/2_lab.docx")
    tex = ""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.color.rgb != RGBColor(0,0,0):
                tex += "1" * len(run.text)
            else:
                tex += "0" * len(run.text)

    while len(tex) % 5 != 0:
        tex += "0"

    print(tex)
    print(bytes.fromhex(hex(int(tex, 2))[2:]).decode(encoding="cp1251"))
    print(bytes.fromhex(hex(int(tex, 2))[2:]).decode(encoding="koi8_r"))
    print(bytes.fromhex(hex(int(tex, 2))[2:]).decode(encoding="cp866"))
    print(mtk2.MTK2_decode(tex))

a = int(input("0 - Замаскировать, 1 - Обнаружить : "))
if a == 0:
    cript()
if a == 1:
    encript()
